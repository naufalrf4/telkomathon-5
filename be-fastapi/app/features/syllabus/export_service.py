import io
import logging
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, cast

import mammoth
from docx import Document as WordDocument
from docxtpl import DocxTemplate, RichText
from sqlalchemy import select

from app.exceptions import NotFoundException, ValidationException
from app.features.syllabus.models import GeneratedSyllabus

_TEMPLATES_DIR = Path(__file__).with_name("templates")
_DOCX_TEMPLATE_PATH = _TEMPLATES_DIR / "syllabus_template.docx"
_PLACEHOLDER_TEXTS = ("Belum tersedia",)
logger = logging.getLogger(__name__)


class SyllabusExportService:
    def __init__(self, db: Any) -> None:
        self.db = db

    async def generate_docx(
        self,
        syllabus_id: uuid.UUID,
        *,
        owner_id: uuid.UUID | None = None,
    ) -> bytes:
        syllabus = await self._get_syllabus(syllabus_id, owner_id=owner_id)
        self._validate_exportable_syllabus(syllabus)
        context = self._build_export_context(syllabus)
        return self._render_docx(context)

    async def generate_pdf(
        self,
        syllabus_id: uuid.UUID,
        *,
        owner_id: uuid.UUID | None = None,
    ) -> bytes:
        syllabus = await self._get_syllabus(syllabus_id, owner_id=owner_id)
        self._validate_exportable_syllabus(syllabus)
        docx_bytes = self._render_docx(self._build_export_context(syllabus))
        return self._render_pdf_from_docx(docx_bytes)

    def _validate_exportable_syllabus(self, syllabus: GeneratedSyllabus) -> None:
        missing_fields = [
            field_name
            for field_name, value in (
                ("tlo", syllabus.tlo),
                ("performance_result", syllabus.performance_result),
                ("condition_result", syllabus.condition_result),
                ("standard_result", syllabus.standard_result),
            )
            if not str(value or "").strip()
        ]
        if missing_fields:
            raise ValidationException(
                "Syllabus export requires complete finalized fields: " + ", ".join(missing_fields)
            )

        journey = syllabus.journey if isinstance(syllabus.journey, dict) else {}
        for stage_name in ("pre_learning", "classroom", "after_learning"):
            stage = journey.get(stage_name)
            if not isinstance(stage, dict):
                raise ValidationException(
                    f"Syllabus export requires a complete journey stage: {stage_name}"
                )
            if not str(stage.get("duration", "")).strip():
                raise ValidationException(
                    f"Syllabus export requires a journey duration for {stage_name}"
                )
            if not str(stage.get("description", "")).strip():
                raise ValidationException(
                    f"Syllabus export requires a journey description for {stage_name}"
                )
            method = stage.get("method")
            content = stage.get("content")
            if not isinstance(method, list) or not [item for item in method if str(item).strip()]:
                raise ValidationException(
                    f"Syllabus export requires journey delivery methods for {stage_name}"
                )
            if not isinstance(content, list) or not [item for item in content if str(item).strip()]:
                raise ValidationException(
                    f"Syllabus export requires journey content for {stage_name}"
                )

    async def _get_syllabus(
        self,
        syllabus_id: uuid.UUID,
        *,
        owner_id: uuid.UUID | None = None,
    ) -> GeneratedSyllabus:
        query = select(GeneratedSyllabus).where(GeneratedSyllabus.id == syllabus_id)
        if owner_id is not None:
            query = query.where(GeneratedSyllabus.owner_id == owner_id)
        result = await self.db.execute(query)
        syllabus = result.scalar_one_or_none()
        if syllabus is None:
            raise NotFoundException("Syllabus", str(syllabus_id))
        return cast(GeneratedSyllabus, syllabus)

    def _build_export_context(self, syllabus: GeneratedSyllabus) -> dict[str, object]:
        journey = syllabus.journey if isinstance(syllabus.journey, dict) else {}
        pre_learning = self._build_stage_context("Pra-pembelajaran", journey.get("pre_learning"))
        classroom = self._build_stage_context("Di kelas", journey.get("classroom"))
        after_learning = self._build_stage_context(
            "Pasca-pembelajaran", journey.get("after_learning")
        )
        elo_items = [item for item in syllabus.elos if isinstance(item, dict)]
        revision_history = (
            [entry for entry in syllabus.revision_history if isinstance(entry, dict)][-3:]
            if syllabus.revision_history
            else []
        )
        course_name = syllabus.course_title or syllabus.topic
        company_profile_summary = syllabus.company_profile_summary or ""
        commercial_overview = syllabus.commercial_overview or ""
        performance_result = syllabus.performance_result or ""
        condition_result = syllabus.condition_result or ""
        standard_result = syllabus.standard_result or ""
        return {
            "date_stamp": self._format_date(syllabus.updated_at or syllabus.created_at),
            "course_title": course_name,
            "course_name": course_name,
            "topic": syllabus.topic,
            "target_level": syllabus.target_level,
            "course_category": syllabus.course_category or "",
            "client_company_name": syllabus.client_company_name or "",
            "tlo": syllabus.tlo,
            "elos": elo_items,
            "journey_stages": [pre_learning, classroom, after_learning],
            "elo_results": self._rich_text(self._join_elo_items(elo_items)),
            "learning_journey_category": self._rich_text(
                syllabus.course_category or "Formal learning"
            ),
            "pre_learning_duration": self._rich_text(cast(str, pre_learning["duration"])),
            "pre_learning_method": self._rich_text(
                self._join_lines(cast(list[str], pre_learning["method"]))
            ),
            "pre_learning_content_list": self._rich_text(
                self._join_lines(cast(list[str], pre_learning["content"]))
            ),
            "pre_learning_evaluation": self._rich_text(cast(str, pre_learning["description"])),
            "classroom_duration": self._rich_text(cast(str, classroom["duration"])),
            "classroom_method": self._rich_text(
                self._join_lines(cast(list[str], classroom["method"]))
            ),
            "classroom_content_list": self._rich_text(
                self._join_lines(cast(list[str], classroom["content"]))
            ),
            "classroom_evaluation": self._rich_text(cast(str, classroom["description"])),
            "after_learning_duration": self._rich_text(cast(str, after_learning["duration"])),
            "after_learning_method": self._rich_text(
                self._join_lines(cast(list[str], after_learning["method"]))
            ),
            "after_learning_content_list": self._rich_text(
                self._join_lines(cast(list[str], after_learning["content"]))
            ),
            "after_learning_evaluation": self._rich_text(cast(str, after_learning["description"])),
            "tlo_result": self._rich_text(syllabus.tlo),
            "performance_result": self._rich_text(performance_result),
            "condition_result": self._rich_text(condition_result),
            "standard_result": self._rich_text(standard_result),
            "company_profile_summary": self._rich_text(company_profile_summary),
            "commercial_overview": self._rich_text(commercial_overview),
            "revision_history": [
                {
                    "summary": str(entry.get("summary", "")).strip() or "Perubahan tanpa ringkasan",
                    "reason": str(entry.get("reason", "")).strip(),
                    "applied_fields": [
                        str(item)
                        for item in cast(list[object], entry.get("applied_fields", []))
                        if isinstance(item, str)
                    ],
                }
                for entry in revision_history
            ],
        }

    def _build_stage_context(self, label: str, value: object) -> dict[str, object]:
        stage = value if isinstance(value, dict) else {}
        method_value = stage.get("method") if isinstance(stage, dict) else []
        content_value = stage.get("content") if isinstance(stage, dict) else []
        return {
            "label": label,
            "duration": str(stage.get("duration", "")).strip() if isinstance(stage, dict) else "",
            "description": str(stage.get("description", "")).strip()
            if isinstance(stage, dict)
            else "",
            "method": [
                str(item).strip()
                for item in method_value
                if isinstance(item, str) and str(item).strip()
            ]
            if isinstance(method_value, list)
            else [],
            "content": [
                str(item).strip()
                for item in content_value
                if isinstance(item, str) and str(item).strip()
            ]
            if isinstance(content_value, list)
            else [],
        }

    def _render_docx(self, context: dict[str, object]) -> bytes:
        if not _DOCX_TEMPLATE_PATH.exists():
            raise FileNotFoundError(f"Syllabus DOCX template not found: {_DOCX_TEMPLATE_PATH}")
        return self._render_docx_from_template(_DOCX_TEMPLATE_PATH, context)

    def _render_docx_from_template(self, template_path: Path, context: dict[str, object]) -> bytes:
        document = DocxTemplate(str(template_path))
        document.render(context)
        buffer = io.BytesIO()
        document.save(buffer)
        return self._strip_placeholder_text(buffer.getvalue())

    def _strip_placeholder_text(self, docx_bytes: bytes) -> bytes:
        document = WordDocument(io.BytesIO(docx_bytes))
        for paragraph in document.paragraphs:
            self._strip_placeholder_from_paragraph(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._strip_placeholder_from_paragraph(paragraph)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _strip_placeholder_from_paragraph(self, paragraph: Any) -> None:
        for run in paragraph.runs:
            updated_text = run.text
            for placeholder in _PLACEHOLDER_TEXTS:
                updated_text = updated_text.replace(placeholder, "")
            if updated_text != run.text:
                run.text = updated_text

    def _render_pdf_from_docx(self, docx_bytes: bytes) -> bytes:
        html_body = self._render_docx_html(docx_bytes)
        try:
            from weasyprint import HTML
        except OSError:
            return self._render_basic_pdf_from_docx(docx_bytes)

        try:
            return cast(bytes, HTML(string=self._wrap_export_html(html_body)).write_pdf())
        except Exception:
            logger.warning(
                "Falling back to basic PDF export because DOCX-to-PDF rendering failed",
                exc_info=True,
            )
            return self._render_basic_pdf_from_docx(docx_bytes)

    def _render_docx_html(self, docx_bytes: bytes) -> str:
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        return str(result.value)

    def _render_basic_pdf_from_docx(self, docx_bytes: bytes) -> bytes:
        lines = self._docx_lines(docx_bytes)
        content_commands = ["BT", "/F1 11 Tf", "50 790 Td"]
        for line in lines:
            content_commands.append(f"({self._escape_pdf_text(line)}) Tj")
            content_commands.append("0 -16 Td")
        content_commands.append("ET")
        content = "\n".join(content_commands).encode("latin-1", errors="replace")

        objects = [
            b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
            b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
            b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
            f"5 0 obj << /Length {len(content)} >> stream\n".encode("latin-1")
            + content
            + b"\nendstream\nendobj\n",
        ]

        pdf = bytearray(b"%PDF-1.4\n")
        offsets: list[int] = []
        for obj in objects:
            offsets.append(len(pdf))
            pdf.extend(obj)

        xref_position = len(pdf)
        pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
        pdf.extend(b"0000000000 65535 f \n")
        for offset in offsets:
            pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
        pdf.extend(
            f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF".encode(
                "latin-1"
            )
        )
        return bytes(pdf)

    def _docx_lines(self, docx_bytes: bytes) -> list[str]:
        document = WordDocument(io.BytesIO(docx_bytes))
        lines = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        return lines[:40]

    def _join_lines(self, items: list[str]) -> str:
        return "\n".join(item for item in items if item.strip()) or "Belum tersedia"

    def _join_elo_items(self, items: list[dict[str, object]]) -> str:
        values = [
            str(item.get("elo", "")).strip() for item in items if str(item.get("elo", "")).strip()
        ]
        return "\n".join(f"• {value}" for value in values) or "Belum tersedia"

    def _rich_text(self, value: str) -> RichText:
        rich = RichText()
        lines = [line.strip() for line in value.splitlines()] or [value]
        wrote_any = False
        for line in lines:
            if not line:
                continue
            if wrote_any:
                rich.add("\n")
            rich.add(line)
            wrote_any = True
        if not wrote_any:
            rich.add("Belum tersedia")
        return rich

    def _wrap_export_html(self, body: str) -> str:
        return f"""
<!DOCTYPE html>
<html lang=\"id\">
  <head>
    <meta charset=\"UTF-8\" />
    <style>
      @page {{ size: A4; margin: 24mm 18mm; }}
      body {{ font-family: sans-serif; color: #172033; font-size: 11pt; line-height: 1.55; }}
      h1, h2, h3, h4 {{ color: #172033; margin: 0 0 10px; }}
      p {{ margin: 0 0 10px; }}
      ul, ol {{ margin: 0 0 10px 20px; }}
      table {{ width: 100%; border-collapse: collapse; margin: 0 0 12px; }}
      td, th {{ border: 1px solid #e2e8f0; padding: 8px; vertical-align: top; }}
    </style>
  </head>
  <body>{body}</body>
</html>
"""

    def _escape_pdf_text(self, value: str) -> str:
        return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    def _format_date(self, value: object) -> str:
        if isinstance(value, datetime):
            return value.strftime("%d %B %Y")
        return ""
