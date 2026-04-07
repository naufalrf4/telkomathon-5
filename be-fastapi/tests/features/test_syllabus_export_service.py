from datetime import UTC, datetime
from io import BytesIO
from typing import Any
from uuid import uuid4

import mammoth
import pytest
from docx import Document as WordDocument

from app.exceptions import ValidationException
from app.features.syllabus.export_service import SyllabusExportService
from app.features.syllabus.models import GeneratedSyllabus


class FakeScalarResult:
    def __init__(self, value: object) -> None:
        self.value = value

    def scalar_one_or_none(self) -> object:
        return self.value


class FakeExportSession:
    def __init__(self, syllabus: GeneratedSyllabus) -> None:
        self.syllabus = syllabus

    async def execute(self, statement: object) -> FakeScalarResult:
        _ = statement
        return FakeScalarResult(self.syllabus)


def build_syllabus() -> GeneratedSyllabus:
    return GeneratedSyllabus(
        id=uuid4(),
        topic="Data Analytics",
        target_level=3,
        course_category="Technical",
        client_company_name="PT Demo",
        course_title="Data Analytics Bootcamp",
        company_profile_summary="Ringkasan perusahaan",
        commercial_overview="Program akselerasi analitik.",
        tlo="Peserta mampu membaca dan menjelaskan data operasional.",
        performance_result="Menyusun insight dari dashboard.",
        condition_result="Dengan studi kasus perusahaan.",
        standard_result="Akurat dan mudah dipahami.",
        elos=[{"elo": "Mengenali metrik utama"}, {"elo": "Menyusun insight"}],
        journey={
            "pre_learning": {
                "duration": "30 menit",
                "method": ["Belajar mandiri"],
                "description": "Pra-belajar",
                "content": ["Artikel pengantar"],
            },
            "classroom": {
                "duration": "1 hari",
                "method": ["Workshop"],
                "description": "Kelas",
                "content": ["Latihan dashboard"],
            },
            "after_learning": {
                "duration": "1 minggu",
                "method": ["Praktik"],
                "description": "Tindak lanjut",
                "content": ["Refleksi hasil"],
            },
        },
        source_doc_ids=[str(uuid4())],
        revision_history=[],
        status="finalized",
        created_at=datetime.now(UTC),
        updated_at=datetime.now(UTC),
    )


def read_document_text(document: Any) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    ]
    return "\n".join([*paragraphs, *table_cells])


@pytest.mark.asyncio
async def test_generate_docx_returns_readable_document() -> None:
    syllabus = build_syllabus()
    service = SyllabusExportService(FakeExportSession(syllabus))

    payload = await service.generate_docx(syllabus.id)

    assert payload[:2] == b"PK"
    document = WordDocument(BytesIO(payload))
    text = read_document_text(document)
    assert "Data Analytics Bootcamp" in text
    assert "Ringkasan perusahaan" in text
    assert "Program akselerasi analitik." in text
    assert "Mengenali metrik utama" in text


@pytest.mark.asyncio
async def test_generate_pdf_returns_pdf_bytes() -> None:
    syllabus = build_syllabus()
    service = SyllabusExportService(FakeExportSession(syllabus))

    payload = await service.generate_pdf(syllabus.id)

    assert payload.startswith(b"%PDF")


@pytest.mark.asyncio
async def test_generate_pdf_uses_rendered_docx_template(monkeypatch: pytest.MonkeyPatch) -> None:
    syllabus = build_syllabus()
    service = SyllabusExportService(FakeExportSession(syllabus))
    await service.generate_docx(syllabus.id)
    captured_docx: dict[str, bytes] = {}

    def fake_convert_to_html(file_obj: BytesIO) -> object:
        captured_docx["payload"] = file_obj.read()

        class Result:
            value = "<p>Rendered from DOCX template</p>"

        return Result()

    class FakeHTML:
        def __init__(self, string: str) -> None:
            self.string = string

        def write_pdf(self) -> bytes:
            assert "Rendered from DOCX template" in self.string
            return b"%PDF-template-preview"

    import sys
    import types

    monkeypatch.setattr(mammoth, "convert_to_html", fake_convert_to_html)
    monkeypatch.setitem(sys.modules, "weasyprint", types.SimpleNamespace(HTML=FakeHTML))

    payload = await service.generate_pdf(syllabus.id)

    assert payload == b"%PDF-template-preview"
    assert captured_docx["payload"][:2] == b"PK"
    rendered_docx = WordDocument(BytesIO(captured_docx["payload"]))
    rendered_text = read_document_text(rendered_docx)
    assert "Data Analytics Bootcamp" in rendered_text


@pytest.mark.asyncio
async def test_generate_docx_rejects_incomplete_finalized_fields() -> None:
    syllabus = build_syllabus()
    syllabus.condition_result = None
    service = SyllabusExportService(FakeExportSession(syllabus))

    with pytest.raises(ValidationException, match="condition_result"):
        await service.generate_docx(syllabus.id)
