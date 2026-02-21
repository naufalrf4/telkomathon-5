import asyncio
from typing import Any

from app.exceptions import FileParseException


async def parse_docx(file_path: str) -> tuple[str, dict[str, Any]]:
    try:
        text, metadata = await asyncio.to_thread(_extract_docx_text, file_path)
        return text, metadata
    except Exception as exc:
        raise FileParseException(f"DOCX parsing failed: {exc}") from exc


def _extract_docx_text(file_path: str) -> tuple[str, dict[str, Any]]:
    from docx import Document

    doc = Document(file_path)
    paragraphs: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = ""
        if para.style and para.style.name:
            style_name = para.style.name
        if style_name.startswith("Heading"):
            headings.append(para.text.strip())
        paragraphs.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n\n".join(paragraphs), {"headings": headings}
